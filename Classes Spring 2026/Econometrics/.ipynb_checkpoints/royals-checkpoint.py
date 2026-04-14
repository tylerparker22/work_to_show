# Create tailored resume and cover letter Word documents
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# ---------------------
# Resume Document
# ---------------------
resume = Document()

resume.add_heading("Tyler C. Parker", level=1)
resume.add_paragraph("Cabot, AR | 501-960-6872 | tyler.parker2216@gmail.com")
resume.add_paragraph("Portfolio: https://lnk.bio/tylerparker22")

resume.add_heading("Education", level=2)
resume.add_paragraph().add_run("Arkansas Tech University – Russellville, AR\n").bold = True
resume.add_paragraph("Bachelor of Science in Business Data Analytics | Expected 2026")

resume.add_paragraph().add_run("North Arkansas College – Harrison, AR\n").bold = True
resume.add_paragraph("Associate of Arts | 2024")
resume.add_paragraph("Dean’s List (3.5 GPA): Fall 2022, Fall & Spring 2023, Spring 2024")

resume.add_heading("Technical Skills", level=2)
resume.add_paragraph("Programming: Python, R, SQL, PHP")
resume.add_paragraph("Data Tools: SQLite, Power BI, Tableau, Snowflake (foundational)")
resume.add_paragraph("Machine Learning: Random Forest, clustering, regression, ANOVA, time-series analysis")
resume.add_paragraph("Visualization: Shiny (R), Dash (Python), interactive dashboards")
resume.add_paragraph("Baseball Analytics: Pitch modeling, strike probability prediction, pitch usage trends, hitter heatmaps")

resume.add_heading("Baseball Analytics Experience", level=2)

resume.add_paragraph().add_run("Baseball Analyst – Arkansas Tech University (2024–2025)\n").bold = True
resume.add_paragraph(
    "- Developed pitching performance models using pitch-by-pitch and video data.\n"
    "- Created tracking models analyzing velocity trends and pitch usage.\n"
    "- Delivered data-driven insights to coaching staff to support player development."
)

resume.add_paragraph().add_run("Freelance Baseball Analytics Developer (2025–Present)\n").bold = True
resume.add_paragraph(
    "- Built an interactive baseball analytics dashboard in Python (Dash).\n"
    "- Designed a relational SQLite database for managing large game datasets.\n"
    "- Implemented Random Forest models to predict next-pitch strike probability and pitch type.\n"
    "- Created hitter heatmaps and pitch usage tables for scouting insights.\n"
    "- Enabled real-time model retraining and probabilistic predictions within the dashboard."
)

resume.add_heading("Data & Analytics Experience", level=2)

resume.add_paragraph().add_run("Student Worker – Business Department, Arkansas Tech University (2025–Present)\n").bold = True
resume.add_paragraph(
    "- Designed and deployed interactive Shiny dashboards for enrollment analytics.\n"
    "- Built SQLite databases and automated ETL workflows from Excel sources.\n"
    "- Conducted trend analysis and created visual reports for administrative leadership."
)

resume.add_paragraph().add_run("National Collegiate Data Analytics Competition (2026)\n").bold = True
resume.add_paragraph(
    "- Placed 42nd nationally using R, machine learning, and visualization tools to analyze complex datasets."
)

resume.add_heading("Additional Experience", level=2)
resume.add_paragraph().add_run("University of Arkansas Division of Agriculture – Temporary Technician (2023–2025)\n").bold = True
resume.add_paragraph(
    "- Collected and tracked large experimental datasets across multiple herbicide trials.\n"
    "- Ensured data integrity and consistency across growing seasons."
)

resume.add_heading("References", level=2)
resume.add_paragraph(
    "Katie Hook – Coordinator of Outreach and Student Success\n"
    "LeMoyne Smith School of Business, Arkansas Tech University\n"
    "khook2@atu.edu | 479-968-0233\n\n"
    "Matt Brown – Interim Associate Dean\n"
    "LeMoyne Smith School of Business, Arkansas Tech University\n"
    "hbrown11@atu.edu | 479-968-0630\n\n"
    "David Pumphrey – Assistant Professor of Business Data Analytics\n"
    "LeMoyne Smith School of Business, Arkansas Tech University\n"
    "dpumphrey@atu.edu | 479-880-4226"
)

resume_path = "/mnt/data/Tyler_Parker_Royals_RnD_Resume.docx"
resume.save(resume_path)

# ---------------------
# Cover Letter Document
# ---------------------
cover = Document()

cover.add_heading("Tyler C. Parker", level=1)
cover.add_paragraph("Cabot, AR | 501-960-6872 | tyler.parker2216@gmail.com")

cover.add_paragraph("\nDear Hiring Committee,")

cover.add_paragraph(
    "I am writing to express my interest in the Analyst – Research & Development position "
    "with the Kansas City Royals. As a Business Data Analytics student at Arkansas Tech University "
    "and a dedicated baseball analytics practitioner, I am eager to contribute advanced statistical "
    "analysis and machine learning solutions to support player evaluation and organizational decision-making."
)

cover.add_paragraph(
    "During my time as a Baseball Analyst at Arkansas Tech University, I developed pitching performance "
    "models using pitch-by-pitch and video data, translating quantitative findings into actionable insights "
    "for coaching staff. I designed tracking models to evaluate velocity trends, pitch usage, and situational outcomes."
)

cover.add_paragraph(
    "Additionally, I independently developed a full-stack baseball analytics dashboard in Python using Dash. "
    "This system integrates a relational SQLite database with Random Forest machine learning models to predict "
    "next-pitch strike probability and pitch type. The dashboard enables real-time model retraining, dynamic filtering, "
    "and advanced visualizations to support scouting and development decisions."
)

cover.add_paragraph(
    "I am passionate about leveraging data to create competitive advantages in baseball operations. "
    "I would welcome the opportunity to contribute analytical rigor, creativity, and collaborative energy "
    "to the Royals’ Research & Development team."
)

cover.add_paragraph("\nSincerely,\nTyler Parker")

cover_path = "C:\Resume"
cover.save(cover_path)

resume_path, cover_path